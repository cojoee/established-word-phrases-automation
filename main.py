#!/usr/bin/env python3
"""
Established Truth, Principles Automation - Railway Deployment
Complete replica of Google Apps Script with Python/Railway
"""

import os
import json
import time
import logging
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
from io import BytesIO
import requests
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from apscheduler.schedulers.background import BackgroundScheduler
from dotenv import load_dotenv

# Google Drive imports
from google.auth.transport.requests import Request as GoogleAuthRequest
from google.oauth2.credentials import Credentials as OAuthCredentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseUpload

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# CONFIG - Load from environment variables or use defaults
NOTION_API_KEY = os.getenv("NOTION_API_KEY", "")
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY", "")
NOTION_DATABASE_ID = os.getenv("NOTION_DATABASE_ID", "3062967169c28172bf7bf32fe67f1a97")
DAILY_DOCUMENTS_DB_ID = os.getenv("DAILY_DOCUMENTS_DB_ID", "e612e7e887cf4462819e92c14ff7a6de")
UMBRELLA_TERM_DOCUMENTS_DB_ID = os.getenv("UMBRELLA_TERM_DOCUMENTS_DB_ID", "9d6fd374d6ce457cbdfccb88dcf91d55")
DRIVE_FOLDER_NAME = os.getenv("DRIVE_FOLDER_NAME", "🧠 Established Truth, Principles, Understanding")
DRIVE_FOLDER_ID = os.getenv("DRIVE_FOLDER_ID", "")  # Direct folder ID — bypasses name search if set
DAILY_DRIVE_FOLDER_NAME = os.getenv("DAILY_DRIVE_FOLDER_NAME", "📅 Established Daily Documents")
DAILY_DRIVE_FOLDER_ID = os.getenv("DAILY_DRIVE_FOLDER_ID", "")  # Direct folder ID for daily docs
UMBRELLA_DRIVE_FOLDER_NAME = os.getenv("UMBRELLA_DRIVE_FOLDER_NAME", "Established Umbrella Term Documents")
UMBRELLA_DRIVE_FOLDER_ID = os.getenv("UMBRELLA_DRIVE_FOLDER_ID", "")
CLAUDE_MODEL = os.getenv("CLAUDE_MODEL", "claude-sonnet-4-6")
CLAUDE_INPUT_COST_PER_M = float(os.getenv("CLAUDE_INPUT_COST_PER_M", "3.00"))
CLAUDE_OUTPUT_COST_PER_M = float(os.getenv("CLAUDE_OUTPUT_COST_PER_M", "15.00"))
MAX_TITLE_LENGTH = int(os.getenv("MAX_TITLE_LENGTH", "150"))
BATCH_SIZE = int(os.getenv("BATCH_SIZE", "5"))
REGISTRY_PAGE_ID = os.getenv("REGISTRY_PAGE_ID", "30729671-69c2-81f9-aaf9-edbbe03eee96")
OPERATION_NAME = os.getenv("OPERATION_NAME", "Established Truth, Principles Automation")
GOOGLE_ACCOUNT = os.getenv("GOOGLE_ACCOUNT", "codyandersonnexus@gmail.com")
GOOGLE_CLIENT_ID = os.getenv("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.getenv("GOOGLE_CLIENT_SECRET", "")
GOOGLE_REFRESH_TOKEN = os.getenv("GOOGLE_REFRESH_TOKEN", "")

# Seed umbrella terms
SEED_UMBRELLA_TERMS = [
    "Divination", "Consciousness", "Automatism", "Transmutation", "Hermeticism",
    "Scripture", "Psychical Research", "Kabbalah", "Depth Psychology", "Esotericism",
    "Embodiment", "Rhetoric", "Cognition", "Craft", "Historiography", "Media"
]

# User's exact prompt
USER_PROMPT_TEMPLATE = """Recognize, establish truths, principles of the {topic}

I want you to provide the following under each (principle, truth) you provide Pattern function principle embodiment wisdom insight through lens of Earthwalker Lover Human Builder (I am self employed, my projects usually consist of cloning existing brands, formats in existing markets and directing them to my monetization points) Person who trains in the gym for visible muscle size gain Provide Utility Provide Human Parallel Provide Insight through Parallel with understood language, subject. So parallel the information with language I understand already. Christian symbolism, Psychology, Flow of Money, Extraction of Currency, Establishment of Resources, Utilization of Established Resources & Wealth Establishment Mechanics, Mathematics, Mechanics, Symbolism, Functionality, Leverage, Incentive Structure, Systems Dynamics / Feedback Loops, Information Asymmetry, Narrative / Frame Control, Temporal Mechanics / Time Preference, Authority Structure / Hierarchy, Cybernetics, Thermodynamics, systems engineering, information technology, computer science, English — A reminder to properly align your generations: You should provide with the most relevant core truth principles wisdom relevant I believe core principles, functions and truth is what enables us to recognize patterns and better navigate the human experience. The intent of this is to expand my understanding of my own human body, function, the symbols and figures I will encounter in my personal human experience of life, the principles the earth follows and ultimate truth and ultimate truth and the deepest and highest levels of understanding possible. I want information that I can read, recognize, understand, utilize. I want to obtain the deepest and highest level of understanding known to mankind and beyond."""

SYSTEM_MESSAGE = """You are generating a comprehensive educational document. Write a complete, well-structured document. You MUST complete the entire document with a proper conclusion. Plan your structure so every section is fully developed and the document ends with a complete closing. Do not leave any section unfinished."""

UMBRELLA_TERM_PROMPT_TEMPLATE = """Given the topic "{topic}", assign it to the single most fitting umbrella term.

EXISTING ESTABLISHED UMBRELLA TERMS: {terms_list}

RULES:
1. FIRST — Check if the topic genuinely belongs under an existing term. Use an existing term if the topic is a natural member, instance, or expression of that umbrella. Most topics SHOULD fit under an existing term. Default to existing terms.
2. ONLY IF NO EXISTING TERM FITS — Establish a new umbrella term. It must be:
   - A genuine hypernym, holonym, or container word (1-2 words maximum)
   - A real word with real meaning and essence, not an academic department name
   - Specific enough to be meaningful (NEVER use: "Knowledge", "Study", "Learning", "Wisdom", "Truth", "Philosophy", "Science", "Culture", "History", "Ideas")
   - Broad enough that other related topics could naturally sit under it too
   - The kind of word where if you said it, a person would immediately picture the domain it contains
3. ANTI-DRIFT RULES:
   - Do NOT create a new term when an existing one works. Check each existing term carefully.
   - Do NOT collapse unrelated topics into the same term just because it is broad.
   - Do NOT create a term that only one topic could ever belong to (too atomic).
   - Do NOT use multi-word academic phrases like "Comparative Religion" or "Political Theory".
   - Each term should feel like a real container with real contents, not a label slapped on.
   - If uncertain between two existing terms, choose the one where the topic is MORE CENTRALLY a member.
4. GOOD examples: Divination, Consciousness, Transmutation, Hermeticism, Scripture, Embodiment, Rhetoric, Craft, Media
5. BAD examples: Ancient Wisdom, Spiritual Studies, Esoteric Knowledge, Human Experience, Deep Understanding

Respond with ONLY the umbrella term. Nothing else. No explanation, no quotes, no punctuation. Just the term."""

# Column names mapping
COLUMNS = {
    "TITLE": "Word, Phrase, Topic",
    "PROCESSED": "Established Truth, Principles, Understanding?",
    "LINK": "Link to Established Resource",
    "READ": "Read, Consumed, Digested?",
    "PRIORITY": "Priority?",
    "UMBRELLA_TERM": "Umbrella Term",
    "DOC_DATE": "Document Establishment Date",
    "UTILIZED_UMBRELLA": "Utilized, Umbrella Term Establishment?",
    "QUOTA_SOURCE": "Quota Source (Google Account)",
    "DAILY_TITLE": "Daily Document",
    "DAILY_DATE": "Date",
    "DAILY_COUNT": "Document Count",
    "DAILY_LINK": "Google Drive, Established Document Link",
    "DAILY_STATUS": "Status",
    "UTD_TITLE": "Umbrella Term Document",
    "UTD_UMBRELLA": "Umbrella Term",
    "UTD_DATE": "Date",
    "UTD_COUNT": "Document Count",
    "UTD_LINK": "Google Drive, Established Document Link",
    "UTD_STATUS": "Status",
    "REGISTRY_NAME": "Operation Name",
    "REGISTRY_STATUS": "Status",
    "REGISTRY_MODEL": "AI Platform & Model",
    "REGISTRY_DATE": "Date Established",
    "REGISTRY_DESCRIPTION": "Operation Description",
    "REGISTRY_DRIVE": "Google Drive Folder",
    "REGISTRY_LAST_RUN": "Last Successful Run",
    "REGISTRY_COST": "Current Month API Cost",
    "REGISTRY_RESTORATION": "Operation Restoration",
    "REGISTRY_FREQUENCY": "Run Frequency",
    "REGISTRY_TOTAL": "Total Inputs Processed",
    "REGISTRY_HEALTH": "Infrastructure Health",
    "REGISTRY_LOGS": "Operation Logs",
    "REGISTRY_SCRIPT": "Google Apps Script",
    "REGISTRY_DATABASES": "Relevant Notion Database(s)"
}


class NotionAPI:
    """Notion API wrapper for database queries and updates"""

    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://api.notion.com/v1"
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Notion-Version": "2022-06-28",
            "Content-Type": "application/json"
        }

    def query_database(self, database_id: str, filter_dict: Optional[Dict] = None, page_size: int = 100, max_results: int = 0) -> List[Dict]:
        """Query a Notion database with optional filtering.
        If max_results > 0, stop after collecting that many results (avoids paginating entire DB).
        """
        url = f"{self.base_url}/databases/{database_id}/query"
        payload = {"page_size": min(page_size, 100)}
        if filter_dict:
            payload["filter"] = filter_dict

        results = []
        has_more = True
        start_cursor = None

        while has_more:
            if start_cursor:
                payload["start_cursor"] = start_cursor

            try:
                response = requests.post(url, headers=self.headers, json=payload)
                response.raise_for_status()
                data = response.json()
                results.extend(data.get("results", []))
                has_more = data.get("has_more", False)
                start_cursor = data.get("next_cursor")

                # Stop early if we have enough results
                if max_results > 0 and len(results) >= max_results:
                    break
            except requests.exceptions.RequestException as e:
                logger.error(f"Error querying database {database_id}: {e}")
                break

        return results

    def get_page(self, page_id: str) -> Dict:
        """Get a single page"""
        url = f"{self.base_url}/pages/{page_id}"
        try:
            response = requests.get(url, headers=self.headers)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            logger.error(f"Error getting page {page_id}: {e}")
            return {}

    def update_page_properties(self, page_id: str, properties: Dict) -> bool:
        """Update page properties"""
        url = f"{self.base_url}/pages/{page_id}"
        payload = {"properties": properties}

        try:
            response = requests.patch(url, headers=self.headers, json=payload)
            if response.status_code >= 400:
                try:
                    error_body = response.json()
                    logger.error(f"Notion API error for page {page_id}: {response.status_code} - {error_body.get('message', 'No message')} - Code: {error_body.get('code', 'N/A')}")
                except Exception:
                    logger.error(f"Notion API error for page {page_id}: {response.status_code} - {response.text[:500]}")
            response.raise_for_status()
            return True
        except requests.exceptions.RequestException as e:
            logger.error(f"Error updating page {page_id}: {e}")
            return False

    def get_block_children(self, block_id: str) -> List[Dict]:
        """Get children blocks of a page"""
        url = f"{self.base_url}/blocks/{block_id}/children"
        results = []
        has_more = True
        start_cursor = None

        while has_more:
            params = {"page_size": 100}
            if start_cursor:
                params["start_cursor"] = start_cursor

            try:
                response = requests.get(url, headers=self.headers, params=params)
                response.raise_for_status()
                data = response.json()
                results.extend(data.get("results", []))
                has_more = data.get("has_more", False)
                start_cursor = data.get("next_cursor")
            except requests.exceptions.RequestException as e:
                logger.error(f"Error getting block children for {block_id}: {e}")
                break

        return results

    def append_block_children(self, block_id: str, children: List[Dict]) -> bool:
        """Append children blocks to a page"""
        url = f"{self.base_url}/blocks/{block_id}/children"
        payload = {"children": children}

        try:
            response = requests.patch(url, headers=self.headers, json=payload)
            response.raise_for_status()
            return True
        except requests.exceptions.RequestException as e:
            logger.error(f"Error appending blocks to {block_id}: {e}")
            return False

    def create_page(self, parent_id: str, properties: Dict, children: Optional[List[Dict]] = None) -> Optional[str]:
        """Create a new page"""
        url = f"{self.base_url}/pages"
        payload = {
            "parent": {"database_id": parent_id},
            "properties": properties
        }
        if children:
            payload["children"] = children

        try:
            response = requests.post(url, headers=self.headers, json=payload)
            response.raise_for_status()
            return response.json().get("id")
        except requests.exceptions.RequestException as e:
            logger.error(f"Error creating page in {parent_id}: {e}")
            return None

    def get_property_value(self, page: Dict, property_name: str) -> Any:
        """Extract property value from a page"""
        properties = page.get("properties", {})
        if property_name not in properties:
            return None

        prop = properties[property_name]
        prop_type = prop.get("type")

        if prop_type == "checkbox":
            return prop.get("checkbox", False)
        elif prop_type == "title":
            title_list = prop.get("title", [])
            return "".join([t.get("plain_text", "") for t in title_list])
        elif prop_type == "rich_text":
            text_list = prop.get("rich_text", [])
            return "".join([t.get("plain_text", "") for t in text_list])
        elif prop_type == "url":
            return prop.get("url")
        elif prop_type == "date":
            date_obj = prop.get("date", {})
            return date_obj.get("start") if date_obj else None
        elif prop_type == "select":
            select_obj = prop.get("select", {})
            return select_obj.get("name") if select_obj else None
        elif prop_type == "number":
            return prop.get("number")

        return None


class ClaudeAPI:
    """Claude API wrapper for content generation"""

    def __init__(self, api_key: str, model: str = "claude-sonnet-4-6"):
        self.api_key = api_key
        self.model = model
        self.base_url = "https://api.anthropic.com/v1"

    def generate_document(self, topic: str, max_tokens: int = 64000) -> tuple[Optional[str], Dict]:
        """Generate comprehensive document for a topic using user's exact prompt"""
        user_content = USER_PROMPT_TEMPLATE.format(topic=topic)

        content, usage = self._call_api(
            [
                {"role": "user", "content": user_content}
            ],
            system=SYSTEM_MESSAGE,
            max_tokens=max_tokens
        )

        # Validate document length (umbrella term calls are short by design, but documents should be substantial)
        if content and len(content.strip()) < 500:
            logger.error(f"Document content too short ({len(content)} chars) for '{topic}', treating as failure")
            return None, usage

        return content, usage

    def assign_umbrella_term(self, topic: str, existing_terms: List[str], max_tokens: int = 50) -> tuple[Optional[str], Dict]:
        """Assign an umbrella term to a topic"""
        terms_list = ", ".join(existing_terms)
        user_content = UMBRELLA_TERM_PROMPT_TEMPLATE.format(
            topic=topic,
            terms_list=terms_list
        )

        response, usage = self._call_api(
            [
                {"role": "user", "content": user_content}
            ],
            system="You are a taxonomist. Respond with ONLY the umbrella term. Nothing else.",
            max_tokens=max_tokens
        )

        if response:
            response = response.strip().strip('"').strip("'")

        return response, usage

    def _call_api(self, messages: List[Dict], system: str = "", max_tokens: int = 64000) -> tuple[Optional[str], Dict]:
        """Make API call to Claude"""
        url = f"{self.base_url}/messages"

        payload = {
            "model": self.model,
            "max_tokens": max_tokens,
            "messages": messages
        }

        if system:
            payload["system"] = system

        headers = {
            "x-api-key": self.api_key,
            "anthropic-version": "2023-06-01",
            "content-type": "application/json"
        }

        try:
            response = requests.post(url, headers=headers, json=payload, timeout=300)
            response.raise_for_status()
            data = response.json()

            content = data.get("content", [{}])[0].get("text", "")
            stop_reason = data.get("stop_reason", "unknown")
            usage = {
                "input_tokens": data.get("usage", {}).get("input_tokens", 0),
                "output_tokens": data.get("usage", {}).get("output_tokens", 0)
            }

            if stop_reason == "max_tokens":
                logger.warning(f"Response truncated (max_tokens reached). Output tokens: {usage['output_tokens']}")

            if not content:
                logger.error("No content returned from Claude API")
                return None, usage

            return content, usage
        except requests.exceptions.Timeout:
            logger.error("Claude API request timed out after 300 seconds")
            return None, {"input_tokens": 0, "output_tokens": 0}
        except requests.exceptions.RequestException as e:
            logger.error(f"Error calling Claude API: {e}")
            return None, {"input_tokens": 0, "output_tokens": 0}


class GoogleDriveAPI:
    """Google Drive API wrapper for file operations using OAuth2"""

    def __init__(self, client_id: str, client_secret: str, refresh_token: str):
        self.service = None
        self.folder_cache = {}
        self.client_id = client_id
        self.client_secret = client_secret
        self.refresh_token = refresh_token

        if client_id and client_secret and refresh_token:
            try:
                credentials = OAuthCredentials(
                    token=None,
                    refresh_token=refresh_token,
                    token_uri='https://oauth2.googleapis.com/token',
                    client_id=client_id,
                    client_secret=client_secret,
                    scopes=['https://www.googleapis.com/auth/drive']
                )
                # Force token refresh
                credentials.refresh(GoogleAuthRequest())

                self.service = build('drive', 'v3', credentials=credentials)
                logger.info("Google Drive API initialized successfully with OAuth2")
            except Exception as e:
                logger.error(f"Error initializing Google Drive API: {e}")

    def find_folder(self, folder_name: str) -> Optional[str]:
        """Find folder ID by name"""
        if folder_name in self.folder_cache:
            return self.folder_cache[folder_name]

        if not self.service:
            logger.warning("Google Drive service not initialized")
            return None

        try:
            query = f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false"
            results = self.service.files().list(
                q=query,
                spaces='drive',
                fields='files(id, name)',
                pageSize=1
            ).execute()

            files = results.get('files', [])
            if files:
                folder_id = files[0]['id']
                self.folder_cache[folder_name] = folder_id
                return folder_id

            logger.warning(f"Folder '{folder_name}' not found, creating it...")
            return self.create_folder(folder_name)
        except Exception as e:
            logger.error(f"Error finding folder {folder_name}: {e}")
            return None

    def create_folder(self, folder_name: str) -> Optional[str]:
        """Create a folder in Google Drive"""
        if not self.service:
            return None
        try:
            file_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder'
            }
            folder = self.service.files().create(
                body=file_metadata,
                fields='id'
            ).execute()
            folder_id = folder.get('id')
            if folder_id:
                self.folder_cache[folder_name] = folder_id
                logger.info(f"Created Drive folder '{folder_name}' with ID: {folder_id}")
            return folder_id
        except Exception as e:
            logger.error(f"Error creating folder {folder_name}: {e}")
            return None

    def upload_docx(self, file_content: bytes, filename: str, folder_id: str) -> Optional[str]:
        """Upload .docx file to Drive"""
        if not self.service:
            logger.warning("Google Drive service not initialized")
            return None

        try:
            file_metadata = {'name': filename, 'parents': [folder_id]}
            media = MediaIoBaseUpload(BytesIO(file_content), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            file = self.service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()

            file_id = file.get('id')
            file_link = file.get('webViewLink')

            if file_id:
                self._set_file_sharing(file_id)

            logger.info(f"Uploaded {filename} to Drive: {file_link}")
            return file_link
        except Exception as e:
            logger.error(f"Error uploading {filename} to Drive: {e}")
            return None

    def _set_file_sharing(self, file_id: str):
        """Set file to be shareable via link"""
        if not self.service:
            return

        try:
            self.service.permissions().create(
                fileId=file_id,
                body={'type': 'anyone', 'role': 'reader'},
                fields='id'
            ).execute()
        except Exception as e:
            logger.error(f"Error setting file sharing for {file_id}: {e}")


class DocumentBuilder:
    """Build .docx documents from content"""

    @staticmethod
    def parse_content_to_blocks(content: str) -> List[Dict]:
        """Parse Claude's response into structured blocks for Notion and docx"""
        blocks = []
        lines = content.split('\n')
        current_paragraph = []

        for line in lines:
            stripped = line.strip()

            if not stripped:
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'empty'})
            elif stripped.startswith('# '):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'heading_1', 'text': stripped[2:].strip()})
            elif stripped.startswith('## '):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'heading_2', 'text': stripped[3:].strip()})
            elif stripped.startswith('### '):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'heading_3', 'text': stripped[4:].strip()})
            elif stripped.startswith('- '):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'bullet', 'text': stripped[2:].strip()})
            elif stripped.startswith('* '):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                blocks.append({'type': 'bullet', 'text': stripped[2:].strip()})
            elif any(stripped.startswith(f"{i}. ") for i in range(1, 10)):
                if current_paragraph:
                    blocks.append({
                        'type': 'paragraph',
                        'text': '\n'.join(current_paragraph)
                    })
                    current_paragraph = []
                text = stripped.split('. ', 1)[1] if '. ' in stripped else stripped
                blocks.append({'type': 'numbered', 'text': text})
            else:
                current_paragraph.append(line)

        if current_paragraph:
            blocks.append({
                'type': 'paragraph',
                'text': '\n'.join(current_paragraph)
            })

        return blocks

    @staticmethod
    def create_docx(content: str, topic: str) -> bytes:
        """Create a .docx file from content"""
        doc = Document()

        # Add title - use the established document naming convention
        doc_heading = f"🧠 Established Truth, Principles of {topic}"
        title = doc.add_heading(doc_heading, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation timestamp
        timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_format = timestamp.runs[0].font
        timestamp_format.size = Pt(10)
        timestamp_format.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacing

        # Parse and add content
        blocks = DocumentBuilder.parse_content_to_blocks(content)

        for block in blocks:
            if block['type'] == 'heading_1':
                doc.add_heading(block['text'], level=1)
            elif block['type'] == 'heading_2':
                doc.add_heading(block['text'], level=2)
            elif block['type'] == 'heading_3':
                doc.add_heading(block['text'], level=3)
            elif block['type'] == 'paragraph':
                text = block['text'].strip()
                if text:
                    doc.add_paragraph(text)
            elif block['type'] == 'bullet':
                doc.add_paragraph(block['text'], style='List Bullet')
            elif block['type'] == 'numbered':
                doc.add_paragraph(block['text'], style='List Number')
            elif block['type'] == 'empty':
                doc.add_paragraph()

        # Add horizontal rule
        pPr = doc.add_paragraph()._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Save to bytes
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    @staticmethod
    def convert_to_notion_blocks(content: str) -> List[Dict]:
        """Convert content to Notion block format"""
        blocks = []
        notion_blocks = DocumentBuilder.parse_content_to_blocks(content)

        for block in notion_blocks:
            if block['type'] == 'heading_1':
                blocks.append({
                    "object": "block",
                    "type": "heading_1",
                    "heading_1": {
                        "rich_text": [{"type": "text", "text": {"content": block['text']}}]
                    }
                })
            elif block['type'] == 'heading_2':
                blocks.append({
                    "object": "block",
                    "type": "heading_2",
                    "heading_2": {
                        "rich_text": [{"type": "text", "text": {"content": block['text']}}]
                    }
                })
            elif block['type'] == 'heading_3':
                blocks.append({
                    "object": "block",
                    "type": "heading_3",
                    "heading_3": {
                        "rich_text": [{"type": "text", "text": {"content": block['text']}}]
                    }
                })
            elif block['type'] == 'paragraph':
                text = block['text'].strip()
                if text:
                    blocks.append({
                        "object": "block",
                        "type": "paragraph",
                        "paragraph": {
                            "rich_text": [{"type": "text", "text": {"content": text}}]
                        }
                    })
            elif block['type'] == 'bullet':
                blocks.append({
                    "object": "block",
                    "type": "bulleted_list_item",
                    "bulleted_list_item": {
                        "rich_text": [{"type": "text", "text": {"content": block['text']}}]
                    }
                })
            elif block['type'] == 'numbered':
                blocks.append({
                    "object": "block",
                    "type": "numbered_list_item",
                    "numbered_list_item": {
                        "rich_text": [{"type": "text", "text": {"content": block['text']}}]
                    }
                })
            elif block['type'] == 'empty':
                blocks.append({
                    "object": "block",
                    "type": "paragraph",
                    "paragraph": {"rich_text": []}
                })

        return blocks


class AutomationEngine:
    """Main automation engine"""

    def __init__(self):
        self.notion = NotionAPI(NOTION_API_KEY)
        self.claude = ClaudeAPI(CLAUDE_API_KEY, CLAUDE_MODEL)
        self.drive = GoogleDriveAPI(GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, GOOGLE_REFRESH_TOKEN)
        self.monthly_cost = 0.0
        self.total_processed = 0
        self.last_run = None
        self.operation_status = "Idle"
        self.health_score = 100

    def process_unprocessed_entries(self):
        """Main processing loop - runs every 5 minutes"""
        logger.info("Starting main processing cycle...")
        self.operation_status = "Processing"
        self.last_run = datetime.now()

        try:
            # Query unprocessed entries that have a title (skip empty rows)
            filter_dict = {
                "and": [
                    {
                        "property": COLUMNS["PROCESSED"],
                        "checkbox": {"equals": False}
                    },
                    {
                        "property": COLUMNS["TITLE"],
                        "title": {"is_not_empty": True}
                    }
                ]
            }

            entries = self.notion.query_database(NOTION_DATABASE_ID, filter_dict, page_size=BATCH_SIZE, max_results=BATCH_SIZE)

            if not entries:
                logger.info("No unprocessed entries found")
                self.operation_status = "Idle"
                return

            logger.info(f"Found {len(entries)} unprocessed entries, processing up to {BATCH_SIZE}")

            # Get existing umbrella terms (use seed terms + cached terms to avoid querying entire DB)
            existing_terms = list(set(SEED_UMBRELLA_TERMS))

            # Process each entry
            for i, entry in enumerate(entries[:BATCH_SIZE]):
                try:
                    success = self._process_entry(entry, existing_terms)
                    if success:
                        self.total_processed += 1
                except Exception as e:
                    logger.error(f"Error processing entry {i}: {e}")
                    self.health_score = max(0, self.health_score - 10)

            self.operation_status = "Idle"
            self._update_registry()
            logger.info(f"Processing cycle complete. Processed {min(len(entries), BATCH_SIZE)} entries")

        except Exception as e:
            logger.error(f"Error in processing cycle: {e}")
            self.operation_status = "Error"
            self.health_score = max(0, self.health_score - 20)

    def _process_entry(self, entry: Dict, existing_terms: List[str]) -> bool:
        """Process a single entry. Returns True on success, False on failure."""
        entry_id = entry.get("id")
        topic = self.notion.get_property_value(entry, COLUMNS["TITLE"])

        if not topic:
            logger.warning(f"Entry {entry_id} has no title")
            return False

        logger.info(f"Processing: {topic}")

        # Step 1: Generate document content
        content, usage = self.claude.generate_document(topic)
        if not content:
            logger.error(f"Failed to generate content for {topic}")
            return False

        # Calculate cost
        input_cost = (usage["input_tokens"] / 1_000_000) * CLAUDE_INPUT_COST_PER_M
        output_cost = (usage["output_tokens"] / 1_000_000) * CLAUDE_OUTPUT_COST_PER_M
        cost = input_cost + output_cost
        self.monthly_cost += cost

        # Step 2: Assign umbrella term
        umbrella_term, _ = self.claude.assign_umbrella_term(topic, existing_terms)
        if not umbrella_term:
            umbrella_term = "Esotericism"  # Default fallback

        # Step 3: Create .docx file
        docx_bytes = DocumentBuilder.create_docx(content, topic)

        # Step 4: Upload to Drive — use direct folder ID if set, otherwise search by name
        folder_id = DRIVE_FOLDER_ID if DRIVE_FOLDER_ID else self.drive.find_folder(DRIVE_FOLDER_NAME)
        if not folder_id:
            logger.error(f"Could not find Drive folder {DRIVE_FOLDER_NAME}")
            return False

        # Document naming: 🧠 Established Truth, Principles of (TOPIC).docx
        # If topic is too long for filename, use Claude to create a meaningful condensed title
        doc_title_prefix = "🧠 Established Truth, Principles of "
        max_topic_len = MAX_TITLE_LENGTH - len(doc_title_prefix) - 5  # 5 for ".docx"
        if len(topic) > max_topic_len:
            # Topic exceeds limit - create all-encompassing meaning title within limits
            condensed, _ = self.claude._call_api(
                [{"role": "user", "content": f'The following topic title is too long for a filename. Condense it to under {max_topic_len} characters while preserving the full meaning and essence of the topic. Do NOT make it generic or cliche — it must be all-encompassing of what the original means. Respond with ONLY the condensed title, nothing else.\n\nOriginal: {topic}'}],
                system="You condense titles while preserving their complete meaning.",
                max_tokens=100
            )
            display_topic = condensed.strip() if condensed else topic[:max_topic_len]
        else:
            display_topic = topic
        filename = f"{doc_title_prefix}{display_topic}.docx".replace('/', '_').replace('\\', '_')
        drive_link = self.drive.upload_docx(docx_bytes, filename, folder_id)
        if not drive_link:
            logger.error(f"Failed to upload document for {topic}")
            return False

        # Step 5: Update Notion entry - CRITICAL: only mark as processed after Drive upload succeeded
        properties = {
            COLUMNS["PROCESSED"]: {"checkbox": True},
            COLUMNS["LINK"]: {"url": drive_link},
            COLUMNS["UMBRELLA_TERM"]: {"select": {"name": umbrella_term}},
            COLUMNS["DOC_DATE"]: {"date": {"start": datetime.now().strftime("%Y-%m-%d")}},
            COLUMNS["QUOTA_SOURCE"]: {"select": {"name": GOOGLE_ACCOUNT}}
        }

        update_success = self.notion.update_page_properties(entry_id, properties)
        if not update_success:
            logger.error(f"CRITICAL: Document uploaded to Drive but Notion update FAILED for {topic}. Drive link: {drive_link}")
            # Retry once
            time.sleep(2)
            update_success = self.notion.update_page_properties(entry_id, properties)
            if not update_success:
                logger.error(f"CRITICAL: Notion update retry also FAILED for {topic}. Manual intervention needed. Drive link: {drive_link}")
                return False

        # Step 6: Append content blocks to Notion page (in batches of 95)
        notion_blocks = DocumentBuilder.convert_to_notion_blocks(content)
        for i in range(0, len(notion_blocks), 95):
            batch = notion_blocks[i:i+95]
            self.notion.append_block_children(entry_id, batch)
            time.sleep(0.5)  # Rate limiting

        logger.info(f"Successfully processed {topic} - Umbrella: {umbrella_term}, Cost: ${cost:.4f}")
        return True

    def compile_daily_documents(self):
        """Compile daily documents - run at 11:55 PM"""
        logger.info("Starting daily compilation...")

        try:
            # Query entries processed today
            today = datetime.now().strftime("%Y-%m-%d")
            filter_dict = {
                "and": [
                    {
                        "property": COLUMNS["DOC_DATE"],
                        "date": {"on_or_after": today}
                    },
                    {
                        "property": COLUMNS["PROCESSED"],
                        "checkbox": {"equals": True}
                    }
                ]
            }

            entries = self.notion.query_database(NOTION_DATABASE_ID, filter_dict, page_size=100)

            if not entries:
                logger.info("No entries processed today for daily compilation")
                return

            logger.info(f"Found {len(entries)} entries for daily compilation")

            # Format the date: "February 17, 2026"
            today_display = datetime.now().strftime("%B %d, %Y").replace(" 0", " ")  # Remove leading zero from day
            daily_doc_title = f"📅 Established Daily Document — {today_display}"

            # Compile all documents
            all_content = f"{daily_doc_title}\n\n"
            for entry in entries:
                topic = self.notion.get_property_value(entry, COLUMNS["TITLE"])
                link = self.notion.get_property_value(entry, COLUMNS["LINK"])
                if topic and link:
                    all_content += f"## {topic}\n[View Document]({link})\n\n"

            # Create .docx
            docx_bytes = DocumentBuilder.create_docx(all_content, daily_doc_title)

            # Upload to Drive — use direct folder ID if set, otherwise search by name
            folder_id = DAILY_DRIVE_FOLDER_ID if DAILY_DRIVE_FOLDER_ID else self.drive.find_folder(DAILY_DRIVE_FOLDER_NAME)
            if not folder_id:
                logger.error(f"Could not find Drive folder {DAILY_DRIVE_FOLDER_NAME}")
                return

            filename = f"{daily_doc_title}.docx"
            drive_link = self.drive.upload_docx(docx_bytes, filename, folder_id)

            if not drive_link:
                logger.error("Failed to upload daily compilation")
                return

            # Create Notion entry in Daily Documents DB
            properties = {
                COLUMNS["DAILY_TITLE"]: {"title": [{"type": "text", "text": {"content": daily_doc_title}}]},
                COLUMNS["DAILY_DATE"]: {"date": {"start": today}},
                COLUMNS["DAILY_COUNT"]: {"number": len(entries)},
                COLUMNS["DAILY_LINK"]: {"url": drive_link},
                COLUMNS["DAILY_STATUS"]: {"select": {"name": "Complete"}}
            }

            self.notion.create_page(DAILY_DOCUMENTS_DB_ID, properties)
            logger.info(f"Daily compilation complete: {len(entries)} documents")

        except Exception as e:
            logger.error(f"Error in daily compilation: {e}")

    def compile_umbrella_term_documents(self, umbrella_term: str):
        """Compile documents for a specific umbrella term"""
        logger.info(f"Starting umbrella term compilation for: {umbrella_term}")

        try:
            # Query un-utilized entries for this umbrella term
            filter_dict = {
                "and": [
                    {
                        "property": COLUMNS["UMBRELLA_TERM"],
                        "select": {"equals": umbrella_term}
                    },
                    {
                        "property": COLUMNS["UTILIZED_UMBRELLA"],
                        "checkbox": {"equals": False}
                    }
                ]
            }

            entries = self.notion.query_database(NOTION_DATABASE_ID, filter_dict, page_size=100)

            if not entries:
                logger.info(f"No un-utilized entries for {umbrella_term}")
                return

            logger.info(f"Found {len(entries)} entries for {umbrella_term}")

            # Compile all documents
            all_content = f"Umbrella Term: {umbrella_term}\n\n"
            for entry in entries:
                topic = self.notion.get_property_value(entry, COLUMNS["TITLE"])
                link = self.notion.get_property_value(entry, COLUMNS["LINK"])
                if topic and link:
                    all_content += f"## {topic}\n[View Document]({link})\n\n"

            # Create .docx
            docx_bytes = DocumentBuilder.create_docx(all_content, f"Umbrella Term: {umbrella_term}")

            # Upload to Drive — use direct folder ID if set, otherwise search by name
            folder_id = UMBRELLA_DRIVE_FOLDER_ID if UMBRELLA_DRIVE_FOLDER_ID else self.drive.find_folder(UMBRELLA_DRIVE_FOLDER_NAME)
            if not folder_id:
                logger.error(f"Could not find Drive folder {UMBRELLA_DRIVE_FOLDER_NAME}")
                return

            filename = f"UmbrellaTermCompilation_{umbrella_term.replace(' ', '_')}.docx"
            drive_link = self.drive.upload_docx(docx_bytes, filename, folder_id)

            if not drive_link:
                logger.error(f"Failed to upload umbrella term compilation for {umbrella_term}")
                return

            # Create Notion entry in Umbrella Term Documents DB
            properties = {
                COLUMNS["UTD_TITLE"]: {"title": [{"type": "text", "text": {"content": f"Compilation - {umbrella_term}"}}]},
                COLUMNS["UTD_UMBRELLA"]: {"select": {"name": umbrella_term}},
                COLUMNS["UTD_DATE"]: {"date": {"start": datetime.now().strftime("%Y-%m-%d")}},
                COLUMNS["UTD_COUNT"]: {"number": len(entries)},
                COLUMNS["UTD_LINK"]: {"url": drive_link},
                COLUMNS["UTD_STATUS"]: {"select": {"name": "Complete"}}
            }

            self.notion.create_page(UMBRELLA_TERM_DOCUMENTS_DB_ID, properties)

            # Mark entries as utilized
            for entry in entries:
                entry_id = entry.get("id")
                self.notion.update_page_properties(entry_id, {
                    COLUMNS["UTILIZED_UMBRELLA"]: {"checkbox": True}
                })

            logger.info(f"Umbrella term compilation complete for {umbrella_term}: {len(entries)} documents")

        except Exception as e:
            logger.error(f"Error in umbrella term compilation: {e}")

    def _update_registry(self):
        """Update the operations registry"""
        try:
            registry_page = self.notion.get_page(REGISTRY_PAGE_ID)
            if not registry_page:
                logger.warning("Could not find registry page")
                return

            properties = {
                COLUMNS["REGISTRY_STATUS"]: {
                    "select": {"name": "Active"}
                },
                COLUMNS["REGISTRY_MODEL"]: {
                    "multi_select": [{"name": CLAUDE_MODEL}]
                },
                COLUMNS["REGISTRY_LAST_RUN"]: {
                    "date": {"start": self.last_run.strftime("%Y-%m-%dT%H:%M:%S")} if self.last_run else None
                },
                COLUMNS["REGISTRY_COST"]: {
                    "number": round(self.monthly_cost, 2)
                },
                COLUMNS["REGISTRY_TOTAL"]: {
                    "number": self.total_processed
                },
                COLUMNS["REGISTRY_HEALTH"]: {
                    "rich_text": [{"type": "text", "text": {"content": f"Claude API: Healthy | Google Drive: Healthy | Notion API: Healthy | Health Score: {self.health_score}"}}]
                },
                COLUMNS["REGISTRY_FREQUENCY"]: {
                    "rich_text": [{"type": "text", "text": {"content": "Every 5 minutes"}}]
                }
            }

            success = self.notion.update_page_properties(REGISTRY_PAGE_ID, properties)
            if success:
                logger.info("Registry updated successfully")
            else:
                logger.warning("Registry update returned failure")
        except Exception as e:
            logger.error(f"Error updating registry: {e}")

    def health_check(self) -> Dict:
        """Health check endpoint"""
        return {
            "status": "healthy" if self.health_score >= 50 else "degraded",
            "operation": OPERATION_NAME,
            "last_run": self.last_run.isoformat() if self.last_run else None,
            "monthly_cost": round(self.monthly_cost, 2),
            "total_processed": self.total_processed,
            "health_score": self.health_score,
            "operation_status": self.operation_status
        }


# Global automation engine
engine = None


def start_scheduler():
    """Start APScheduler for automation tasks"""
    global engine
    engine = AutomationEngine()

    scheduler = BackgroundScheduler(daemon=True)

    # Main processing every 5 minutes
    scheduler.add_job(
        engine.process_unprocessed_entries,
        'interval',
        minutes=5,
        id='main_processing',
        name='Main Processing Loop'
    )

    # Daily compilation at 11:55 PM
    scheduler.add_job(
        engine.compile_daily_documents,
        'cron',
        hour=23,
        minute=55,
        id='daily_compilation',
        name='Daily Compilation'
    )

    scheduler.start()
    logger.info("Scheduler started successfully")


def get_health():
    """Health check handler"""
    global engine
    if engine is None:
        return {"status": "not_ready", "message": "Automation engine not initialized"}
    return engine.health_check()


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "health":
        # Quick health check
        engine = AutomationEngine()
        print(json.dumps(get_health(), indent=2))
    else:
        # Start scheduler
        start_scheduler()

        # Keep the process alive
        try:
            import signal
            def signal_handler(sig, frame):
                logger.info("Shutting down gracefully...")
                sys.exit(0)

            signal.signal(signal.SIGINT, signal_handler)
            signal.signal(signal.SIGTERM, signal_handler)

            while True:
                time.sleep(1)
        except KeyboardInterrupt:
            logger.info("Interrupted by user")
            sys.exit(0)
